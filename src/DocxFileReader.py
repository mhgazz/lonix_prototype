'''
Created on Oct 28, 2019

@author: root
'''
from AbstractFileReader import AbstractFileReader
import docx

class DocxFileReader(AbstractFileReader):
    '''
    implementation for txt plain file
    '''

    i = 0

    def __init__(self):
        '''
        Constructor
        '''
    
    def openFile(self, curFile):
        fp = docx.Document(curFile)
        return fp
        
    def readNextSegment(self, fp):
        curParagraph = fp.paragraphs[self.i]
        self.i = self.i + 1
        return curParagraph.text
    
    def readCertSegment(self, fp, order):
        curParagraph = fp.paragraphs[order]
        return curParagraph.text
    
    def closeFile(self,curFile):
        curFile = None
        
    def readFile(self,fp):
        return fp.paragraphs.__len__()
        